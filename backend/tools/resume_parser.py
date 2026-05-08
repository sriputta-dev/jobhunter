"""
Resume Parser
--------------
Extracts plain text from uploaded resume files.
Supports: PDF (.pdf) and Word (.docx) formats.

Used by:
  - /api/resume/upload endpoint
  - AI agents read extracted text instead of hardcoded background

How it works:
  PDF  → pdfminer extracts text layer (works for text-based PDFs)
       → pypdf as fallback
  DOCX → python-docx reads paragraphs and table cells
"""

import io
import logging
from typing import Tuple

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfminer with pypdf fallback."""
    text = ""

    # Try pdfminer first — best quality for text PDFs
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        from pdfminer.layout import LAParams

        laparams = LAParams(line_margin=0.5, word_margin=0.1)
        text = pdfminer_extract(
            io.BytesIO(file_bytes),
            laparams=laparams,
        )
        if text and text.strip():
            logger.info(f"pdfminer extracted {len(text)} chars from PDF")
            return _clean_text(text)
    except Exception as e:
        logger.warning(f"pdfminer failed: {e}")

    # Fallback: pypdf
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text)
        text = "\n".join(pages)
        if text.strip():
            logger.info(f"pypdf extracted {len(text)} chars from PDF")
            return _clean_text(text)
    except Exception as e:
        logger.warning(f"pypdf failed: {e}")

    logger.error("All PDF extraction methods failed")
    return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        parts = []

        # Paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Tables (skills tables, education tables etc)
        for table in doc.tables:
            for row in table.rows:
                row_parts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_parts.append(cell_text)
                if row_parts:
                    parts.append(" | ".join(row_parts))

        text = "\n".join(parts)
        logger.info(f"python-docx extracted {len(text)} chars from DOCX")
        return _clean_text(text)

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def extract_resume_text(filename: str, file_bytes: bytes) -> Tuple[str, str]:
    """
    Main entry point. Detects file type and extracts text.
    Returns (extracted_text, error_message).
    error_message is empty string on success.
    """
    filename_lower = filename.lower().strip()

    if filename_lower.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
        if not text.strip():
            return "", (
                "Could not extract text from this PDF. "
                "Make sure it is a text-based PDF (not a scanned image). "
                "Try saving your resume as DOCX instead."
            )
        return text, ""

    elif filename_lower.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
        if not text.strip():
            return "", "Could not extract text from this DOCX file."
        return text, ""

    elif filename_lower.endswith(".doc"):
        return "", (
            "Old .doc format is not supported. "
            "Please save your resume as .docx (Word 2007 or later) or .pdf and upload again."
        )

    elif filename_lower.endswith(".txt"):
        try:
            text = file_bytes.decode("utf-8", errors="replace")
            return _clean_text(text), ""
        except Exception:
            return "", "Could not read text file."

    else:
        ext = filename.rsplit(".", 1)[-1] if "." in filename else "unknown"
        return "", (
            f"Unsupported file type: .{ext}. "
            "Please upload a .pdf or .docx file."
        )


def _clean_text(text: str) -> str:
    """Clean extracted text — remove excessive whitespace, fix encoding."""
    import re
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse 3+ blank lines into 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove null bytes
    text = text.replace("\x00", "")
    return text.strip()
